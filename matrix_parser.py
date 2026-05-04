#!/usr/bin/env python3
"""Парсер матрицы переходов из PDF / DOCX / XLSX / CSV / JSON.

Каноничный формат:
{
    "columns": [{"id": 58, "name": "Foydali 45"}, ...],   # целевые тарифы (НА что переходим)
    "rows":    [{"id": 84, "name": "Bor 90+90",
                 "expected": {"58": "+", "46": "+", ...}}, ...]  # исходные
}

Парсинг построен максимально терпимо: вытаскивает текст -> ищет ID в круглых скобках
вида "(ID 58)" или "ID: 58" -> значения "+"/"-" (а также "yes/no/✓/✗/✅/❌/да/нет").
"""
from __future__ import annotations

import io
import re
import json
import unicodedata
from typing import Any


_ID_RE = re.compile(r"\bID\s*[:#№]?\s*(\d+)\b", re.IGNORECASE)
# Запасной маркер: "(58)" в самом конце строки — некоторые матрицы пишут только число в скобках
_TAIL_PAREN_NUM_RE = re.compile(r"\(\s*(\d{1,5})\s*\)\s*$")
_NUMBER_RE = re.compile(r"\b(\d{2,5})\b")


# Все варианты «плюсоподобных» символов — что в матрице может встретиться как «доступно».
_PLUS_CHARS = {
    "+",        # ASCII PLUS
    "＋",   # ＋ FULLWIDTH PLUS SIGN
    "➕",   # ➕ HEAVY PLUS SIGN
    "✓",   # ✓ CHECK MARK
    "✔",   # ✔ HEAVY CHECK MARK
    "✅",   # ✅ WHITE HEAVY CHECK MARK
    "☑",   # ☑ BALLOT BOX WITH CHECK
}

# Все варианты «минусоподобных» символов — что может встретиться как «недоступно».
_MINUS_CHARS = {
    "-",        # ASCII HYPHEN-MINUS
    "−",   # − MINUS SIGN
    "–",   # – EN DASH
    "—",   # — EM DASH
    "‐",   # ‐ HYPHEN
    "‑",   # ‑ NON-BREAKING HYPHEN
    "‒",   # ‒ FIGURE DASH
    "﹣",   # ﹣ SMALL HYPHEN-MINUS
    "－",   # － FULLWIDTH HYPHEN-MINUS
    "✕",   # ✕ MULTIPLICATION X
    "✖",   # ✖ HEAVY MULTIPLICATION X
    "✗",   # ✗ BALLOT X
    "✘",   # ✘ HEAVY BALLOT X
    "❌",   # ❌ CROSS MARK
    "×",   # × MULTIPLICATION SIGN
}

# Слова целиком (после lower) — однозначное «да/нет».
_PLUS_WORDS = {
    "yes", "true", "available", "да", "ha",
    "доступно", "доступен", "доступна",
    "разрешено", "включено",
}
_MINUS_WORDS = {
    "no", "false", "blocked", "нет", "yo'q", "yoq", "yoʻq",
    "недоступно", "недоступен", "недоступна",
    "запрещено", "отключено", "закрыт",
}

# Regex для очистки ячеек: символы нулевой ширины и NBSP-варианты пробелов.
_ZW_RE   = re.compile("[​‌‍⁠﻿­]")
_NBSP_RE = re.compile("[     　 ᠎ -  ]")


def _normalize_cell(c) -> str:
    """Привести значение ячейки к str с одинарными пробелами.

    Сначала вырезаем символы нулевой ширины (ZWSP, BOM, ZWJ, soft hyphen),
    затем NBSP и его родственников (U+00A0, U+202F, U+2007, U+3000, ...) меняем
    на обычный пробел. Иначе сравнение и поиск маркеров ID ломаются.
    """
    if c is None:
        return ""
    s = str(c)
    s = _ZW_RE.sub("", s)
    s = _NBSP_RE.sub(" ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _has_id_marker(s: str) -> bool:
    """В ячейке есть распознаваемый маркер ID тарифа?"""
    if not s:
        return False
    return bool(_ID_RE.search(s)) or bool(_TAIL_PAREN_NUM_RE.search(s))


def _norm_value(v: Any) -> str:
    """Привести значение ячейки к '+' / '-' / '' (пусто).

    Принципы:
    - «+» ставим только когда уверены: ячейка состоит из плюс-символов
      (+ ＋ ➕ ✓ ✔ ✅ ☑) или является ключевым словом (yes/да/available/ha/доступно…).
    - «-» ставим только при однозначном минус-символе (- − – — ‐ ‑ ﹣ － ✕ ✖ ✗ ✘ ❌ ×)
      или ключевом слове (no/нет/yo'q/blocked/недоступно…).
    - Excel-флаг «1»/«0» интерпретируем только если ячейка состоит ровно из этой цифры.
    - «Y»/«N» больше НЕ считаем флагами — случайная буква в ячейке раньше давала сбои.
    - Всё остальное (нераспознанный текст) → пусто. Никакого фолбэка в «-» (именно это
      приводило к ошибкам, когда необычный текст тихо превращался в «-»).
    """
    if v is None:
        return ""

    # NFKC: fullwidth/совместимые символы → канонические формы
    s = unicodedata.normalize("NFKC", str(v))
    s = _ZW_RE.sub("", s)
    s = _NBSP_RE.sub(" ", s)
    s = s.strip()
    if not s:
        return ""

    # 1) Одиночный символ — самый частый случай
    if len(s) == 1:
        if s in _PLUS_CHARS:
            return "+"
        if s in _MINUS_CHARS:
            return "-"

    # 2) Ячейка состоит только из плюс- (или минус-) символов (например "++", "——")
    no_space = re.sub(r"\s+", "", s)
    if no_space and all(ch in _PLUS_CHARS for ch in no_space):
        return "+"
    if no_space and all(ch in _MINUS_CHARS for ch in no_space):
        return "-"

    # 3) Целое слово (yes/no/да/нет/доступно…)
    s_low = s.lower()
    if s_low in _PLUS_WORDS:
        return "+"
    if s_low in _MINUS_WORDS:
        return "-"

    # 4) Excel-флаг — только ровно "1" или "0"
    if s_low == "1":
        return "+"
    if s_low == "0":
        return "-"

    # 5) Не распознали — оставляем пусто (в отчёте это станет статусом "info")
    return ""


def _extract_id(text: str) -> int | None:
    """Найти ID тарифа в строке вида 'Foydali 45 (ID 58)'.

    Приоритет: маркер 'ID 58' > последнее число в скобках в конце.
    """
    m = _ID_RE.search(text)
    if m:
        return int(m.group(1))
    m = _TAIL_PAREN_NUM_RE.search(text)
    if m:
        return int(m.group(1))
    return None


def _name_only(text: str) -> str:
    """Убрать '(ID 58)' и ведущие/хвостовые пробелы."""
    s = re.sub(r"\(ID[^)]*\)", "", text, flags=re.IGNORECASE)
    s = re.sub(r"\s+", " ", s).strip(" \t\n\r-—()")
    return s


def _build(rows_raw: list[list]) -> dict:
    """Из произвольной таблицы собрать каноничный JSON.

    Многоуровневые шапки (например, верхняя строка-«Переход НА тарифный план»
    объединена через мерж) и пустые строки автоматически пропускаются:
    мы ищем строку с максимальным количеством маркеров `(ID …)` —
    это и есть заголовок. Всё, что выше — игнорируется. Всё, что ниже —
    обрабатывается как данные (но только те строки, где cell[0] тоже содержит
    маркер ID — это настоящие исходные тарифы).
    """
    if not rows_raw:
        raise ValueError("Таблица пуста")

    # Нормализация: убираем пробелы/переносы, NBSP → пробел.
    normalized: list[list[str]] = []
    for r in rows_raw:
        row = [_normalize_cell(c) for c in (r or [])]
        if any(row):
            normalized.append(row)
    if not normalized:
        raise ValueError("Таблица пуста (после очистки от пустых строк)")

    # 1) Найти строку-заголовок: больше всего ячеек с ID-маркером в [1:].
    header_idx = -1
    best_count = 1  # требуем минимум 2 колонки с ID
    for i, row in enumerate(normalized[:12]):  # шапка обычно в первых строках
        cnt = sum(1 for c in row[1:] if _has_id_marker(c))
        if cnt > best_count:
            best_count = cnt
            header_idx = i
    if header_idx < 0:
        raise ValueError(
            "Не найдена строка-заголовок: ни в одной из первых 12 строк нет двух и более "
            "ячеек с маркером 'ID …' (или '… (58)' в конце). Убедитесь, что в шапке у "
            "тарифов указан ID."
        )

    header = normalized[header_idx]

    # 2) Собрать колонки — только те, у которых есть ID. Помним их позицию,
    # чтобы корректно прочитать ячейки в строках данных (там могут быть
    # пустые «merge-shadow» столбцы).
    real_cols: list[tuple[int, dict]] = []
    for orig_idx, h in enumerate(header[1:], start=1):
        if not h or not _has_id_marker(h):
            continue
        col_id = _extract_id(h)
        if col_id is None:
            continue
        real_cols.append((orig_idx, {"id": col_id, "name": _name_only(h), "raw": h}))

    if len(real_cols) < 2:
        raise ValueError(
            f"В заголовке найдено всего {len(real_cols)} колонок с ID. "
            "Проверьте, что у тарифов указан ID в формате '(ID 58)'."
        )

    columns_clean = [c for _, c in real_cols]

    # 3) Строки данных — после header_idx. Берём только те, где cell[0]
    # содержит маркер ID (исходный тариф).
    rows: list[dict] = []
    for r in normalized[header_idx + 1:]:
        if not r or not r[0]:
            continue
        if not _has_id_marker(r[0]):
            continue
        row_id = _extract_id(r[0])
        if row_id is None:
            continue
        expected: dict[str, str] = {}
        for orig_idx, col in real_cols:
            cell_val = r[orig_idx] if orig_idx < len(r) else ""
            expected[str(col["id"])] = _norm_value(cell_val)
        rows.append({
            "id": row_id, "name": _name_only(r[0]),
            "raw": r[0], "expected": expected,
        })

    if not rows:
        raise ValueError(
            "Не найдено ни одной строки с исходным тарифом (с маркером 'ID …' в первой колонке)."
        )

    return {"columns": columns_clean, "rows": rows}


# ======================== FORMAT-SPECIFIC PARSERS ========================

def parse_json(content: bytes) -> dict:
    data = json.loads(content.decode("utf-8"))
    # если уже каноничный — вернуть как есть
    if isinstance(data, dict) and "columns" in data and "rows" in data:
        return data
    raise ValueError("JSON должен содержать поля columns и rows")


def parse_csv(content: bytes) -> dict:
    import csv
    text = content.decode("utf-8-sig", errors="replace")
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(io.StringIO(text), dialect)
    rows = [r for r in reader if any((c or "").strip() for c in r)]
    if not rows:
        raise ValueError("CSV пуст")
    return _build(rows)


def parse_xlsx(content: bytes) -> dict:
    """Excel с поддержкой merged cells: значение из левой-верхней ячейки мержа
    разливается на все остальные ячейки этого мержа — иначе шапка теряется.
    """
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active

    # Раскладываем merge-cells: тиражируем значение в каждую ячейку диапазона
    merged_map: dict[tuple[int, int], str] = {}
    for rng in list(ws.merged_cells.ranges):
        top_left = ws.cell(row=rng.min_row, column=rng.min_col).value
        if top_left is None:
            continue
        for r in range(rng.min_row, rng.max_row + 1):
            for c in range(rng.min_col, rng.max_col + 1):
                merged_map[(r, c)] = str(top_left)

    rows: list[list[str]] = []
    for row_idx, row in enumerate(ws.iter_rows(values_only=False), start=1):
        out_row = []
        for col_idx, cell in enumerate(row, start=1):
            v = merged_map.get((row_idx, col_idx))
            if v is None:
                v = cell.value
            out_row.append("" if v is None else str(v))
        if any(c.strip() for c in out_row):
            rows.append(out_row)
    if not rows:
        raise ValueError("XLSX пуст")
    return _build(rows)


def parse_pdf(content: bytes) -> dict:
    """Извлечь таблицу из PDF через pdfplumber. Объединяет таблицы со
    всех страниц (некоторые PDF разбивают шапку и первые строки на отдельные
    страницы) — _build сам разберётся, где заголовок."""
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        raise ValueError(
            "Для PDF нужен пакет `pdfplumber` (pip install pdfplumber). "
            "Либо приложите файл в формате XLSX / CSV / JSON."
        )

    all_tables: list[list[list[str]]] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            for t in tables:
                cleaned = [[("" if c is None else str(c)) for c in row] for row in t if row]
                if len(cleaned) >= 1 and any(any(c.strip() for c in r) for r in cleaned):
                    all_tables.append(cleaned)
    if not all_tables:
        raise ValueError("В PDF не найдено читаемой таблицы (pdfplumber)")

    # 1) Берём самую большую таблицу — обычно матрица из 1 куска.
    biggest = max(all_tables, key=lambda t: sum(len(r) for r in t))
    try:
        return _build(biggest)
    except ValueError as e_one:
        # 2) Если не вышло — склеиваем все таблицы (PDF мог разбить страницы).
        merged: list[list[str]] = []
        for t in all_tables:
            merged.extend(t)
        try:
            return _build(merged)
        except ValueError:
            raise e_one  # выбрасываем первую (более информативную) ошибку


def parse_docx(content: bytes) -> dict:
    """Извлечь самую большую таблицу из .docx с разворотом merge-ячеек."""
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise ValueError("Для DOCX нужен пакет `python-docx` (pip install python-docx)") from e
    doc = Document(io.BytesIO(content))
    if not doc.tables:
        raise ValueError("В DOCX нет таблиц")
    table = max(doc.tables, key=lambda t: sum(len(r.cells) for r in t.rows))
    rows = []
    for row in table.rows:
        # python-docx уже отдаёт «развёрнутые» merged-ячейки как повторяющиеся
        # значения в каждой ячейке мержа.
        rows.append([cell.text.strip() for cell in row.cells])
    rows = [r for r in rows if any(c for c in r)]
    if len(rows) < 2:
        raise ValueError("Таблица в DOCX слишком маленькая")
    return _build(rows)


# ======================== PUBLIC API ========================

def parse(filename: str, content: bytes) -> dict:
    name = (filename or "").lower()
    if name.endswith(".json"):
        return parse_json(content)
    if name.endswith(".csv") or name.endswith(".tsv"):
        return parse_csv(content)
    if name.endswith(".xlsx") or name.endswith(".xlsm"):
        return parse_xlsx(content)
    if name.endswith(".pdf"):
        return parse_pdf(content)
    if name.endswith(".docx"):
        return parse_docx(content)
    raise ValueError(f"Неподдерживаемый формат: {filename}. Используйте PDF/DOCX/XLSX/CSV/JSON.")


def stats(spec: dict) -> dict:
    rows = spec.get("rows", [])
    cols = spec.get("columns", [])
    plus = minus = empty = 0
    for r in rows:
        for v in (r.get("expected") or {}).values():
            if v == "+":
                plus += 1
            elif v == "-":
                minus += 1
            else:
                empty += 1
    return {
        "rows": len(rows),
        "columns": len(cols),
        "cells": len(rows) * len(cols),
        "expected_plus": plus,
        "expected_minus": minus,
        "expected_empty": empty,
    }
