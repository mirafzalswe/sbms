#!/usr/bin/env python3
"""Парсер списка тарифов для теста доступности продукта.

В отличие от matrix_parser (двумерная таблица «откуда → куда»), здесь нужен
плоский список тарифных планов: [{id, name}].

Поддерживаются три способа ввода:
1. Файл (XLSX / CSV / JSON / PDF / DOCX) — любая таблица, в которой есть колонка
   с маркером "(ID 58)" / "ID: 58" / "(58)" в конце ячейки.
2. Текст (paste) — каждая строка трактуется как один тариф; подходит для копирования
   из Excel/Word.
3. JSON — массив `[{id, name}]` или просто `[123, 456, 789]`.

Утилиты извлечения ID и имён переиспользуем из matrix_parser, чтобы поведение
было идентичным основной матричной странице.
"""
from __future__ import annotations

import io
import json
import re
from typing import Any

from matrix_parser import (
    _ID_RE,
    _TAIL_PAREN_NUM_RE,
    _extract_id,
    _has_id_marker,
    _name_only,
    _normalize_cell,
)


def _dedup(items: list[dict]) -> list[dict]:
    """Удалить дубли по id, сохранив порядок и более полное имя."""
    seen: dict[int, dict] = {}
    for it in items:
        rid = it.get("id")
        if rid is None:
            continue
        try:
            rid = int(rid)
        except (TypeError, ValueError):
            continue
        if rid in seen:
            # если у уже сохранённого имя пустое — обновим
            if not seen[rid].get("name") and it.get("name"):
                seen[rid]["name"] = it["name"]
            continue
        seen[rid] = {"id": rid, "name": (it.get("name") or "").strip()}
    return list(seen.values())


def _from_table_rows(rows: list[list[str]]) -> list[dict]:
    """Из произвольной таблицы вытащить тарифы.

    Стратегия: пробегаем все ячейки; если в ячейке есть ID-маркер — считаем
    её записью тарифа. Имя берём из этой же ячейки (через _name_only). Этого
    достаточно для одностолбцовых списков и не ломается, если в файле
    несколько столбцов.
    """
    out: list[dict] = []
    for row in rows or []:
        for cell in row or []:
            s = _normalize_cell(cell)
            if not s or not _has_id_marker(s):
                continue
            rid = _extract_id(s)
            if rid is None:
                continue
            out.append({"id": rid, "name": _name_only(s)})
    return _dedup(out)


# ======================== FORMAT-SPECIFIC PARSERS ========================

def parse_text(text: str) -> list[dict]:
    """Парсит paste-текст. Каждая непустая строка обрабатывается отдельно.

    Также поддерживается формат «id, name» / «id name» / просто число — для
    случаев, когда в строке нет маркера "(ID …)". Если ни один из этих форматов
    не сработал — строка пропускается.
    """
    out: list[dict] = []
    for line in (text or "").splitlines():
        s = _normalize_cell(line)
        if not s:
            continue
        # 1) "Foydali 45 (ID 58)" — стандартный маркер
        rid = _extract_id(s)
        if rid is not None:
            out.append({"id": rid, "name": _name_only(s)})
            continue
        # 2) "58, Foydali 45" / "58 Foydali 45" / "58;Foydali 45" / "58\tFoydali 45"
        m = re.match(r"^\s*(\d{1,6})\s*[,;:\t\s\-—–|]+\s*(.+)$", s)
        if m:
            try:
                out.append({"id": int(m.group(1)), "name": m.group(2).strip()})
                continue
            except ValueError:
                pass
        # 3) "58" — просто число, без имени
        if re.fullmatch(r"\d{1,6}", s):
            out.append({"id": int(s), "name": ""})
    return _dedup(out)


def parse_json(content: bytes | str) -> list[dict]:
    """JSON: массив [{id, name}] / [{ratePlanId, name}] / [123, 456] / {tariffs: [...]}."""
    raw = content.decode("utf-8") if isinstance(content, (bytes, bytearray)) else content
    data = json.loads(raw)
    if isinstance(data, dict):
        # обёртка вида {tariffs: [...]} или {items: [...]}
        for key in ("tariffs", "items", "rateplans", "ratePlans", "data"):
            if isinstance(data.get(key), list):
                data = data[key]
                break
        else:
            raise ValueError("В JSON не найдено поля tariffs/items/data со списком тарифов")
    if not isinstance(data, list):
        raise ValueError("JSON должен содержать массив или объект с массивом тарифов")

    out: list[dict] = []
    for it in data:
        if isinstance(it, (int, float)):
            out.append({"id": int(it), "name": ""})
            continue
        if isinstance(it, str):
            # строка вида "Foydali 45 (ID 58)"
            rid = _extract_id(it)
            if rid is not None:
                out.append({"id": rid, "name": _name_only(it)})
                continue
            if re.fullmatch(r"\d{1,6}", it.strip()):
                out.append({"id": int(it.strip()), "name": ""})
                continue
            continue
        if isinstance(it, dict):
            # ищем id под разными именами
            rid = it.get("id") or it.get("ratePlanId") or it.get("rateplan_id")
            if rid is None:
                # вложенная структура, как в API SBMS: {ratePlan: {ratePlanId, name}}
                rp = it.get("ratePlan") if isinstance(it.get("ratePlan"), dict) else None
                if rp:
                    rid = rp.get("ratePlanId")
                    name = rp.get("name") or ""
                else:
                    name = it.get("name") or ""
            else:
                name = it.get("name") or ""
            try:
                out.append({"id": int(rid), "name": str(name).strip()})
            except (TypeError, ValueError):
                continue
    return _dedup(out)


def parse_csv(content: bytes) -> list[dict]:
    import csv
    text = content.decode("utf-8-sig", errors="replace")
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(io.StringIO(text), dialect)
    rows = [r for r in reader if any((c or "").strip() for c in r)]
    if not rows:
        raise ValueError("CSV пуст")
    items = _from_table_rows(rows)
    if not items:
        # fallback: возможно это просто столбец вида "id,name" без ID-маркера
        items = parse_text(text)
    if not items:
        raise ValueError("Не найдено ни одной записи с ID тарифа")
    return items


def parse_xlsx(content: bytes) -> list[dict]:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active
    rows: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        out_row = ["" if v is None else str(v) for v in row]
        if any(c.strip() for c in out_row):
            rows.append(out_row)
    if not rows:
        raise ValueError("XLSX пуст")
    items = _from_table_rows(rows)
    if not items:
        raise ValueError(
            "Не найдено ни одной ячейки с маркером 'ID …' / '(58)'. "
            "Убедитесь, что у тарифов в таблице указан ID."
        )
    return items


def parse_pdf(content: bytes) -> list[dict]:
    try:
        import pdfplumber  # type: ignore
    except ImportError as e:
        raise ValueError(
            "Для PDF нужен пакет `pdfplumber` (pip install pdfplumber). "
            "Либо приложите файл в формате XLSX / CSV / JSON / TXT."
        ) from e

    all_rows: list[list[str]] = []
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for t in (page.extract_tables() or []):
                for r in t:
                    if not r:
                        continue
                    all_rows.append(["" if c is None else str(c) for c in r])
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                pass

    items = _from_table_rows(all_rows)
    if not items:
        items = parse_text("\n".join(text_parts))
    if not items:
        raise ValueError("В PDF не найдено читаемой таблицы или строк с ID тарифов")
    return items


def parse_docx(content: bytes) -> list[dict]:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise ValueError("Для DOCX нужен пакет `python-docx` (pip install python-docx)") from e
    doc = Document(io.BytesIO(content))
    rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
    items = _from_table_rows(rows)
    if not items:
        # пробуем плоский текст параграфов
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        items = parse_text(text)
    if not items:
        raise ValueError("В DOCX не найдено таблицы или строк с ID тарифов")
    return items


# ======================== PUBLIC API ========================

def parse_file(filename: str, content: bytes) -> list[dict]:
    name = (filename or "").lower()
    if name.endswith(".json"):
        return parse_json(content)
    if name.endswith(".csv") or name.endswith(".tsv"):
        return parse_csv(content)
    if name.endswith(".xlsx") or name.endswith(".xlsm"):
        return parse_xlsx(content)
    if name.endswith(".pdf"):
        return parse_pdf(content)
    if name.endswith(".docx"):
        return parse_docx(content)
    if name.endswith(".txt"):
        return parse_text(content.decode("utf-8", errors="replace"))
    raise ValueError(
        f"Неподдерживаемый формат: {filename}. Используйте PDF/DOCX/XLSX/CSV/JSON/TXT."
    )


def parse(*, filename: str | None = None, content: bytes | None = None,
          text: str | None = None, json_data: Any | None = None) -> list[dict]:
    """Универсальная точка входа.

    Вызывающий код передаёт ровно один из вариантов:
    - filename + content — файл;
    - text — paste-текст;
    - json_data — уже распарсенный JSON (объект/массив).
    """
    if json_data is not None:
        return parse_json(json.dumps(json_data, ensure_ascii=False))
    if text is not None:
        items = parse_text(text)
        if not items:
            raise ValueError("Не удалось распознать ни одного тарифа в тексте")
        return items
    if content is not None:
        return parse_file(filename or "", content)
    raise ValueError("Нужен один из параметров: filename+content, text или json_data")
